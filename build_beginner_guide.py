from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Quran_Kalima_Beginner_Code_Guide.docx'
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)

def font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic

def shade(cell, fill):
    node = OxmlElement('w:shd')
    node.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(node)

def para(doc, text='', label=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if label and text.startswith(label):
        font(p.add_run(label), bold=True)
        font(p.add_run(text[len(label):]))
    else:
        font(p.add_run(text))
    return p

def h(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), {1:16, 2:13, 3:12}[level], BLUE if level < 3 else DARK, True)
    return p

def bullet(doc, text, numbered=False):
    p = doc.add_paragraph(style='List Number' if numbered else 'List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))

def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, label, width in zip(t.rows[0].cells, headers, widths):
        cell.width = Inches(width)
        shade(cell, 'E8EEF5')
        font(cell.paragraphs[0].add_run(label), 9.5, DARK, True)
    for row in rows:
        cells = t.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            cell.width = Inches(width)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.08
            font(cell.paragraphs[0].add_run(value), 9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def box(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.cell(0,0)
    shade(c, 'F4F6F9')
    font(c.paragraphs[0].add_run(title), 10.5, DARK, True)
    p = c.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    font(p.add_run(body), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def f(doc, path, purpose, works, change, careful=''):
    h(doc, path, 3)
    para(doc, 'Simple meaning: ' + purpose, 'Simple meaning: ')
    para(doc, 'How it works: ' + works, 'How it works: ')
    para(doc, 'Change here when: ' + change, 'Change here when: ')
    if careful:
        para(doc, 'Be careful: ' + careful, 'Be careful: ')

doc = Document()
s = doc.sections[0]
s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
s.header_distance = s.footer_distance = Inches(.492)
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)
for n, size, color in [('Heading 1',16,BLUE),('Heading 2',13,BLUE),('Heading 3',12,DARK)]:
    styles[n].font.name = 'Calibri'
    styles[n].font.size = Pt(size)
    styles[n].font.color.rgb = color
    styles[n].font.bold = True
header = s.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run('Quran Kalima - Beginner Code Guide'), 9, RGBColor(90,90,90))
footer = s.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run('Beginner-friendly guide | Read before asking AI to change code'), 9, RGBColor(90,90,90))

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80); font(p.add_run('Quran Kalima'),28,DARK,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Beginner Code Guide'),22,BLUE,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(28); font(p.add_run('Roman Urdu / simple English mein: file ka kaam, code ka flow, aur AI prompt kaise likhein.'),12,RGBColor(90,90,90),False,True)
box(doc, 'Is guide ka goal', 'Aapko programmer banne ki zarurat nahi hai. Iske baad aap feature ko file se connect kar paayenge, safe change request likh paayenge, aur AI ke answer ko check kar paayenge.')
para(doc, 'Project folder: C:\\Users\\sharafatulla.khan\\Desktop\\quran_vocab', 'Project folder: ')
para(doc, 'Important rule: Pehle feature identify karein, phir related file dekhein, phir AI ko exact file name aur expected result bataein. Bina samjhe multiple files delete/replace na karein.', 'Important rule: ')
doc.add_page_break()

h(doc, 'Folder Guide')
for name, meaning in [
('lib/','Aapke app ka main Dart code. Normally yahin feature change hoga.'),
('lib/screens/','Full pages aur bottom sheets. UI/interaction change ka first stop.'),
('lib/services/','Data aur business logic. Save/load/API/SRS/Firebase related feature yahan dekhein.'),
('lib/providers/','Dark mode, display settings aur login user jaisi shared state.'),
('lib/models/','Word/Surah/Ayah ka data shape. New data field yahan se start hota hai.'),
('lib/widgets/','Repeat hone wale small UI components. Same word UI kai jagah update karna ho to yahan change karein.'),
('lib/data/','Hard-coded/static Quran support data, jaise ruku endings.'),
('assets/data/','JSON/text: translations, word meanings, morphology tags. Edit before backup.'),
('assets/fonts/','Arabic/Urdu font files. New font ke liye pubspec.yaml bhi update hoga.'),
('pubspec.yaml','Packages, assets aur font registration. New package/font/image ke liye important.'),
('firestore.rules','Firebase database security. Galat change se user data expose ho sakta hai.'),
('test/','Automated tests. Abhi actual app coverage bahut kam hai.'),
]: para(doc, f'{name} - {meaning}', f'{name} - ')

h(doc, 'Core files - detail mein')
f(doc, 'lib/main.dart', 'Yeh app ka main switchboard hai. Flutter yahin se start hota hai.', 'Firebase initialize hota hai; theme/morphology/translation/glossary load hote hain; providers create hote hain; auth ke basis par login ya main tabs open hote hain.', 'App title, startup loading, initial screen, provider ya startup order badalna ho.', 'Wrong change se poora app start nahi hoga. Startup mein heavy cheezein already load ho rahi hain; bina plan aur add na karein.')
f(doc, 'lib/screens/main_navigation.dart', 'Bottom navigation ka owner: Quran, Vocabulary, Progress, Profile tabs.', 'IndexedStack se chaar pages memory mein preserve hote hain; selected index badalne par tab switch hota hai.', 'Tab name/icon/order change, new fifth tab add, ya initial tab change karna ho.', 'New tab ke liye screen import, `_screens` list aur NavigationDestination tino update hone chahiye.')
f(doc, 'lib/screens/auth_screen.dart', 'Login/Register/Google sign-in ka page.', 'Text controllers email/password/name ko hold karte hain. Firebase Auth sign-in ke baad main.dart auth stream screen switch karta hai. Guest button directly main navigation kholta hai.', 'Login design, validation, password reset, Google button text, registration fields add karna ho.', 'Guest mode current root auth state se independent hai; is flow ko carefully improve karna chahiye.')
f(doc, 'lib/providers/user_provider.dart', 'Logged-in user aur profile data ka shared manager.', 'Firebase auth changes listen karta hai; Firestore se profile load karta hai; sync restore run karta hai; profile update aur sign out handle karta hai.', 'User name, daily goal, photo, gender ya profile rules change karne ho.', 'Cloud restore/sync changes ke saath sync_service.dart bhi samajhna zaroori hai.')
f(doc, 'lib/providers/theme_provider.dart', 'Dark/light mode aur ek Arabic font preference ka manager.', 'SharedPreferences se `is_dark_mode` aur `arabic_font` read/write karta hai. Quran list/reader/profile isko listen karte hain.', 'Dark mode colors ya global theme badalna ho.', 'Arabic font setting display_provider.dart mein bhi hai; dono ko ek jaisa rakhna important hai.')
f(doc, 'lib/providers/display_provider.dart', 'Reader display settings: Arabic/Urdu size, font, line height, spacing, grammar colours.', 'Constructor preferences se load karta hai; setters save karke UI listeners ko notify karte hain.', 'Word tile ka size/font/colour/spacing change karna ho.', 'Isme loaded state nahi hai; defaults pehle dikhen aur baad mein preferences apply ho sakti hain.')

h(doc, 'Quran reading files')
f(doc, 'lib/screens/surah_list_screen.dart', '114 Surah ki list, search icon, progress, bookmarks aur flashcard entry button.', 'Startup par local known-word progress, last read aur bookmarks load karta hai. Card tap SurahReader kholta hai. Search delegate ko callback deta hai.', 'Surah card design, progress badge, app bar, bookmark section ya flashcard button badalna ho.', 'Yeh file lambi hai; card code `_SurahCard` class mein hai, main page code top mein hai.')
f(doc, 'lib/widgets/surah_search_delegate.dart', 'Search popup jo Surah name/number ya `2:255` input handle karta hai.', 'Query se 114 Surah scan karta hai. Valid surah:ayah par reader ko jump location callback deta hai.', 'Search placeholder, name matching, Juz search ya results UI change karna ho.', 'Text/translation word search yahan implemented nahi hai; iske liye alag index/service lagega.')
f(doc, 'lib/screens/surah_reader_screen.dart', 'App ka sabse important Quran reading page.', 'Selected Surah ka verse count quran package se leta hai; MorphologyService word objects banata hai; WordTile render hota hai. Scroll last-read save karta hai; bookmark, ruku/juz, translations aur settings control karta hai.', 'Reader layout, show/hide translation, bookmark UI, word click flow, mushaf mode ya reader controls change karna ho.', 'Yeh 1,000+ lines ka file hai. Translation/network logic modify karte waqt duplicate calls ka dhyan rakhein.')
f(doc, 'lib/widgets/word_tile.dart', 'Reader mein har Arabic word aur uska meaning dikhane wala reusable piece.', 'QuranWord leta hai, display provider se font/colour leta hai, morphology segments se grammar colour lagata hai. Tap/long press callbacks reader se milte hain.', 'Har word ka Arabic style, known word mein meaning hide, grammar colours, Urdu/English meaning layout change karna ho.', 'Yeh change reader ke saare words par lagega; isi liye reader file mein duplicate styling na karein.')
f(doc, 'lib/widgets/word_detail_dialog.dart', 'Word tap karne par khulne wala popup.', 'Arabic word, meaning, audio URL, known toggle, copy/share aur Grammar button dikhata hai. Iska own audio player aur animations hain.', 'Popup design, button labels, share format, audio button ya known-state UI change karna ho.', 'Known toggle ke baad reader callback update hota hai; is connection ko break na karein.')
f(doc, 'lib/screens/morphology_sheet.dart', 'Grammar & Morphology bottom sheet.', 'Selected word ke segments se Sarf chain/root forms banata hai. User ayah ke doosre word ko tap kar sakta hai; root examples aur word audio dikhata hai.', 'Grammar tab design, English/Urdu explanation, root usage list, audio UI change karna ho.', 'Actual grammar data morphology_service.dart se aata hai. Is file mein shared audio subscription lifecycle issue ho sakta hai.')

h(doc, 'Vocabulary, flashcards aur progress files')
f(doc, 'lib/screens/vocabulary_screen.dart', 'All/Known/Unknown words ki searchable list aur swipe actions.', 'Known words WordProgressService se aate hain; list frequency, Arabic/Urdu search aur occurrence screen navigation use karti hai.', 'Vocabulary tab design, sorting, filters, swipe wording ya empty screen change karna ho.', 'Current code mein `getWordFrequenciesFromGlossary()` call hota hai, lekin service method empty/incorrect type hai. Isko fix kiye bina feature reliable nahi hoga.')
f(doc, 'lib/screens/word_occurrences_screen.dart', 'Ek word Quran mein kin ayahs mein aata hai woh dikhata hai.', 'Word ke matching ayahs load karta hai, target word highlight karta hai, aur selected occurrence se reader ko jump karwata hai.', 'Occurrence card UI, highlight colour, result count, reader jump behavior change karna ho.', 'Full results data/index par depend karta hai; locally discovered data se list incomplete ho sakti hai.')
f(doc, 'lib/screens/flashcard_screen.dart', 'SRS flashcard learning page.', 'SrsService se due/new cards leta hai; card flip, known/unknown/delete, points, sample ayah, audio, root/morphology aur session summary handle karta hai. `FlashWord` local data holder hai.', 'Flashcard design, buttons, flip/swipe, session message, points display change karna ho.', 'SRS rule change karna ho to saath srs_service.dart bhi change hoga. UI file bahut lambi hai; business rule yahan hard-code na karein.')
f(doc, 'lib/screens/progress_screen.dart', 'Learning dashboard: overall percent, streak, heatmap, surah progress, badges.', 'SharedPreferences aur WordProgressService se numbers load karta hai; animation controllers se ring/bars animate karta hai.', 'Progress cards, motivational text, badges, chart layout ya calculation display change karna ho.', 'Screen sirf initState par reload hota hai; tab switch ke baad latest values visible na hon to refresh logic add karna padega.')
f(doc, 'lib/screens/profile_settings_screen.dart', 'Profile aur settings ka page.', 'UserProvider, ThemeProvider, DisplayProvider use karta hai; display preview, font/size sliders, sync status, profile fields, account actions aur external links dikhata hai.', 'Theme/display control, daily goal, profile UI, logout/delete account, support links change karna ho.', 'Reminder abhi sirf time picker/snackbar hai; notification schedule nahi hota. Share bhi placeholder snackbar hai.')

h(doc, 'Data models aur static data files')
f(doc, 'lib/models/word.dart', '`QuranWord` ek word ka complete UI data holder hai: Arabic, meaning, known status, segments, root/lemma/POS.', 'Reader MorphologyService se QuranWord banata hai; WordTile/Dialog/Morphology use karte hain.', 'Naya word field jaise note, favorite flag, extra translation add karna ho.', 'Field add karne ke baad object banane wali jagah aur `copyWith` update karni hogi.')
f(doc, 'lib/models/surah.dart', '`Surah` Surah id, names aur verse count hold karta hai.', 'Surah list/search/occurrence reader ko Surah object pass karte hain.', 'Surah mein new display data add karna ho.', 'Current urduName actual Urdu source se fill nahi ho raha; name behavior check karein.')
f(doc, 'lib/models/ayah.dart', '`Ayah` verse number aur words ka simple model hai.', 'Abhi active app flow mein meaningful use nahi dikh raha.', 'Future mein structured verse state banana ho.', 'Unused code ho sakta hai; bina use samjhe delete na karein.')
f(doc, 'lib/data/surah_data.dart', '`buildSurahList()` quran package se 114 Surah model list banata hai.', 'Surah metadata convert karta hai.', 'Ek single central Surah list chahiye.', 'List screen apni metadata generation bhi karta hai; duplicate approach hai.')
f(doc, 'lib/data/ruku_data.dart', 'Har Surah ke Ruku end ayah numbers ka static map.', 'Reader is map se Ruku badge show karta hai.', 'Ruku boundaries ya label logic change karna ho.', 'Data accuracy critical hai; source verify kiye bina values edit na karein.')

doc.add_page_break()
h(doc, 'Services: app ka background engine')
f(doc, 'lib/services/word_progress_service.dart', 'Known/unknown word progress ka local database helper.', 'Arabic harakat remove karke key banata hai; SharedPreferences mein known words, daily counts, Surah word lists/counts aur saved meanings store karta hai. Sync schedule karta hai.', 'Known button behavior, normalization, total progress, daily count ya frequency storage change karna ho.', 'Core data file hai. Key name badalne par old users ka data missing lagega. Migration plan ke bina key rename na karein.')
f(doc, 'lib/services/srs_service.dart', 'Flashcards ka scheduling engine.', 'Har word ke liye `srs_<word>` preference record; known par stage/points/due session advance, unknown par stage reduce; session list save/restore.', 'Review interval, points, daily goal logic, due-card priority change karna ho.', 'Current system calendar dates nahi, app sessions use karta hai. Daily behaviour improve karna ho to carefully redesign karein.')
f(doc, 'lib/services/sync_service.dart', 'Local progress ko Firestore cloud se sync/restore karta hai.', 'Known words, SRS cards, daily stats, Surah data preferences se read karke Firestore docs mein write karta hai. Login par cloud restore karta hai.', 'Sync status, cloud data fields, debounce time, conflict handling change karna ho.', 'High-risk file. Multiple accounts same phone par local preferences mix ho sakte hain; Firestore document size limits bhi concern hain.')
f(doc, 'lib/services/morphology_service.dart', 'Arabic grammar corpus ka brain.', 'Text asset parse karke WordSegment map mein store karta hai; root, lemma, POS, Sarf chain aur reader words build karta hai.', 'Grammar labels/explanations, morphology parsing, word-segment display data change karna ho.', 'Large/complex service hai. Corpus format/parsing change se reader/grammar dono break honge. AI ko sample input-output dein.')
f(doc, 'lib/services/word_glossary_service.dart', 'Word-by-word Urdu/English/Hindi meaning loader.', 'JSON asset selected language ke liye memory cache mein load karta hai. Position key `surah:ayah:word` se meaning deta hai.', 'Word meaning language, new glossary JSON, HTML cleaning, language labels change karna ho.', 'New language ke liye asset file, pubspec.yaml, glossary map aur UI option sab add honge.')
f(doc, 'lib/services/translation_service.dart', 'Ayah translation loader.', 'Selected scholar ka bundled asset memory mein load karta hai. Agar translation memory/prefs mein nahi hai to AlQuran Cloud HTTP API try karta hai.', 'New scholar, translation asset, translation network cache, RTL handling change karna ho.', 'Only one active bundled scholar. Network error silent null ho sakta hai; user-friendly error add karna useful hai.')
f(doc, 'lib/services/quran_cache_service.dart', 'Quran ayahs aur word locations ka cache/index helper.', 'quran package se fast first-word location index build karta hai; optional CDN full ayah cache functions bhi hain.', 'Word occurrence search/index performance, offline ayah cache change karna ho.', '`initialize()` ka current main flow mein caller nahi dikh raha. Cache complete flag failed downloads ke baad bhi set ho sakta hai.')
f(doc, 'lib/services/quran_preloader_service.dart', 'All 114 Surah ke word counts/preload helper.', 'quran package verses se words count karta hai aur progress service mein save karta hai.', 'Full vocabulary database pehle se build karna ho.', 'Current flow mein use nahi ho raha; commented old code bhi hai.')
f(doc, 'lib/services/corpus_service.dart', 'Optional online Quran corpus helper.', 'Quran.com API se morphology/frequency try karta hai aur audio/corpus URL banata hai.', 'Remote corpus integration ya frequency use karna ho.', 'Current screens isko use nahi kar rahe; before adding, rate limits/error/cache design karein.')

h(doc, 'Assets aur configuration')
f(doc, 'pubspec.yaml', 'Flutter project configuration.', 'Packages, assets folders aur custom fonts declare karta hai. `flutter pub get` isi file ke baad packages download/resolve karta hai.', 'New package, image/font/data file add karna ho.', 'YAML indentation important hai. Package version change ke baad test zaroor karein.')
f(doc, 'assets/data/*.json and quran_morphology.txt', 'Bundled translation, word meaning aur morphology content.', 'Services startup ya language selection par files read karke maps banate hain.', 'Content correction/new translation/new language add karna ho.', 'JSON format break hua to app empty meanings dikhayega ya startup errors hide kar dega. Backup + JSON validation karein.')
f(doc, 'assets/fonts/*', 'Arabic/Urdu display fonts.', 'pubspec font family name se widgets font select karte hain.', 'New font ya font quality issue.', 'Font add ke baad pubspec registration aur app restart/hot restart chahiye.')
f(doc, 'firestore.rules', 'Cloud data permission rules.', 'Only signed-in user ko apne `/users/{uid}` data read/write allow karta hai.', 'New Firestore collection ya backend security policy.', 'Production data security file hai. "allow read, write: if true" kabhi na karein.')
f(doc, 'test/widget_test.dart', 'Automated test folder ka current default test.', 'Flutter test runner isko run karta hai.', 'Har important feature change ke test add karna ho.', 'Abhi project ke real features covered nahi hain; AI se feature change ke saath test bhi maangein.')

h(doc, 'Code ko padhne ka easy formula')
para(doc, 'Kisi screen file ko khol kar is order mein dekhein:', 'Kisi screen file ko khol kar is order mein dekhein:')
for x in [
'Top imports: file kin services/widgets/providers ko use kar rahi hai?',
'Class name: screen ka naam kya hai? StatefulWidget matlab page ke andar values change hoti hain.',
'State variables: underscore `_` se start names private hain; `_isLoading`, `_knownWords` bataate hain page kya remember karta hai.',
'`initState()`: screen open hote hi kya data load hota hai?',
'`build()`: user ko actual UI kya dikhta hai?',
'Methods `onTap`, `_load...`, `_save...`, `_toggle...`: user action/data logic trace karein.',
'`Navigator.push`: naya page khul raha hai.',
]: bullet(doc,x,True)
box(doc, 'Useful words ka simple meaning', '`await` = pehle kaam khatam hone ka wait. `setState` = screen ko dobara draw karo. `Future` = result baad mein aayega. `Provider` = shared setting/state. `SharedPreferences` = phone par chhota local storage. `Firebase/Firestore` = cloud login/database. `static` = class ka shared/global function/data.')

h(doc, 'AI ko strong prompt kaise dein')
para(doc, 'Bad prompt: "Reader improve karo."', 'Bad prompt: ')
para(doc, 'Good prompt mein yeh sab hona chahiye:', 'Good prompt mein yeh sab hona chahiye:')
for x in [
'Goal: user ko exactly kya change dikhe?',
'File: is guide se likely file path mention karein.',
'Current behavior: abhi kya ho raha hai?',
'Expected behavior: change ke baad kya hona chahiye?',
'Keep safe: kaunse existing behavior break nahi hone chahiye?',
'Verification: `flutter analyze` aur related test/run karna hai.',
]: bullet(doc,x)
h(doc, 'Copy-paste prompt templates', 2)
box(doc, 'Reader UI change', 'Open `lib/screens/surah_reader_screen.dart`. I want [exact UI change]. Keep word tap, long press known toggle, bookmarks, translations and current navigation working. First explain which methods/widgets you will modify. Then make the smallest change. Run `flutter analyze` and tell me any errors.')
box(doc, 'Word popup change', 'In `lib/widgets/word_detail_dialog.dart`, change [button/layout/text]. Do not change audio URL format, known-word callback, copy/share, or morphology navigation. Show before/after behavior and test it.')
box(doc, 'Flashcard rule change', 'Review `lib/services/srs_service.dart` and `lib/screens/flashcard_screen.dart`. I want [new rule]. Put scheduling/business rule in SrsService, not only UI. Preserve existing saved user data or provide migration. Add tests for known and unknown review outcomes.')
box(doc, 'New setting', 'Add a [setting name] setting. Use existing pattern in `lib/providers/display_provider.dart` or `theme_provider.dart`: default, SharedPreferences load, setter saves and notifies listeners, UI control in `profile_settings_screen.dart`, then use it in [widget]. Do not duplicate setting in another provider.')
box(doc, 'Bug report', 'I see this issue: [steps]. It happens on [screen]. Trace UI file to service/provider, explain root cause before changing code, propose smallest safe fix, then run analysis/tests.')
box(doc, 'New feature', 'I want [feature]. First list affected files and data model/storage changes. Reuse existing architecture where safe. Do not delete existing code. Implement in small steps, add empty/loading/error state, and test main flow.')
h(doc, 'AI se hamesha yeh 5 questions poochhein', 2)
for x in [
'Is change se exactly kaunsi files modify hongi aur har file kyun?',
'Kya existing saved user data, Firebase sync, bookmarks, SRS ya navigation par impact hoga?',
'Kya code duplicate ho raha hai? Kya existing provider/service reuse ho sakti hai?',
'Failure case kya hai: internet off, empty data, user logged out, old preferences?',
'Mujhe manual test steps do: app mein kya click karke verify karun?',
]: bullet(doc,x,True)

h(doc, 'Important warnings before you change code')
for x in [
'Preference keys (`known_word_`, `srs_`, `surah_`) rename/delete mat karein unless migration plan ho; otherwise old progress gayab lagega.',
'Firebase sync/auth/account delete ko simple UI change na samjhein. Yeh real user data handle karta hai.',
'Morphology JSON/text ka format change karne se reader words/grammar blank ho sakte hain.',
'`pubspec.yaml` edit ke baad `flutter pub get` chalana hota hai.',
'AI ko vague request par multiple files replace karne na dein. Diff/reason pehle dekhein.',
'Current worktree mein already source changes ho sakte hain. Unrelated files overwrite/reset na karwayen.',
]: bullet(doc,x)
h(doc, 'Your practical next steps')
for x in [
'Pehle Vocabulary compile issue fix karwayen: word_progress_service.dart aur vocabulary_screen.dart ka typed, tested fix.',
'Phir reader settings ko one provider mein consolidate karwayen, kyunki font setting abhi do places mein hai.',
'Uske baad progress refresh aur real reminder/share implement karwayen.',
'Firebase sync/account logic ko release se pehle senior Flutter/Firebase review ke through harden karwayen.',
'Har new feature ke saath AI se "affected files + test plan + rollback plan" maangein.',
]: bullet(doc,x,True)
h(doc, 'One-page mental map')
para(doc, 'User sees page -> Screen. Same small UI repeated -> Widget. Data shape -> Model. Save/load/rules -> Service. Global setting/user state -> Provider. App starts -> main.dart.')
box(doc, 'When you are stuck', 'Mujhe feature ka naam, screen ka screenshot/steps, desired result aur is guide ka file path bhej dein. Example: "Vocabulary tab mein search ka result Urdu mein highlight karna hai; likely lib/screens/vocabulary_screen.dart. Please explain first, then change only needed files."')

doc.save(OUT)
print(OUT)

h(doc, 'Start Here: App ko 5 simple parts mein samjhein')
table(doc, ['Part','Aap isko kya samjhein','Folder'], [
('Screens','App ke pages jo user dekhta hai: Quran list, reader, vocabulary, progress, profile.','lib/screens/'),
('Services','Background helpers: data load/save, SRS, Firebase sync, translation, morphology.','lib/services/'),
('Providers','Global settings/state: dark mode, font size, logged-in user.','lib/providers/'),
('Models','Word, Surah aur Ayah data ke fields.','lib/models/'),
('Widgets','Reusable UI pieces: word tile, popup, search.','lib/widgets/'),
], [1.05,3.8,1.65])
box(doc, 'Golden rule', 'Screen tells you what user sees. Service tells you data kaise aata/save hota hai. Provider tells you poore app mein setting/state kaise change hoti hai. Model tells you data mein fields kya hain.')
h(doc, 'Jab user app open karta hai: simple flow')
for x in [
'App `lib/main.dart` se start hota hai.',
'Firebase, preferences, theme, Arabic morphology, translation aur word glossary load hote hain.',
'Logged-in user ko `MainNavigation`; logged-out user ko `AuthScreen` milta hai.',
'MainNavigation mein 4 tabs hain: Quran, Vocabulary, Progress, Profile.',
'Quran tab: SurahListScreen -> SurahReaderScreen -> WordDetailDialog -> MorphologySheet.',
'Known/Unknown word `WordProgressService` phone par save karta hai; login ho to `SyncService` cloud update try karta hai.',
]: bullet(doc,x,True)
h(doc, 'Sabse common change: kaunsi file kholni hai?')
table(doc,['Aap kya change karna chahte hain?','Pehli file','Related file'],[
('Start/login flow','lib/main.dart','auth_screen.dart, user_provider.dart'),
('Bottom tabs','screens/main_navigation.dart','new screen file'),
('Surah list/search','screens/surah_list_screen.dart','widgets/surah_search_delegate.dart'),
('Reader/translation/bookmark','screens/surah_reader_screen.dart','widgets/word_tile.dart, translation_service.dart'),
('Word popup/audio/share','widgets/word_detail_dialog.dart','screens/morphology_sheet.dart'),
('Grammar/root/Sarf','screens/morphology_sheet.dart','services/morphology_service.dart'),
('Vocabulary list/swipe','screens/vocabulary_screen.dart','services/word_progress_service.dart'),
('Flashcard/SRS','screens/flashcard_screen.dart','services/srs_service.dart'),
('Progress/streak','screens/progress_screen.dart','services/word_progress_service.dart'),
('Theme/font/profile','screens/profile_settings_screen.dart','providers/theme_provider.dart, display_provider.dart'),
('Cloud sync','services/sync_service.dart','providers/user_provider.dart, firestore.rules'),
], [2.4,2.55,1.55])
h(doc,'Edit karne se pehle safe process')
for x in [
'Goal likhein: "Reader mein Urdu translation ka default font size 18 chahiye."',
'Is guide ke change-map se main file choose karein.',
'AI ko feature + file + expected UI + kya preserve karna hai + test command batayein.',
'AI se pehle explain karwayen ki files kaunsi change hongi aur kyun.',
'Change ke baad `flutter analyze` aur app run/test karwayen. Exact error AI ko paste karein.',
'Ek time par ek feature change karein; working version ka Git commit ya ZIP backup rakhein.',
]: bullet(doc,x,True)
doc.add_page_break()